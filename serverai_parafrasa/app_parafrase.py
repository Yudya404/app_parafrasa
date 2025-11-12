import os 
import re
import logging
import threading
import uuid
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from docx import Document
import torch
import time

# ===============================
# ⚙️ Konfigurasi dasar
# ===============================
app = Flask(__name__)
CORS(app)
os.makedirs("../logs", exist_ok=True)
os.makedirs("../tmp_paraphrase", exist_ok=True)

logging.basicConfig(
    filename="../logs/app_parafrasa.log",
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
logging.info("🚀 Flask Paraphraser Async dimulai.")

# ===============================
# 🧠 Model utama
# ===============================
MODEL_NAME = "ramsrigouthamg/t5-large-paraphraser-diverse-high-quality"
DEVICE = 0 if torch.cuda.is_available() else -1
device_str = "cuda" if DEVICE == 0 else "cpu"
logging.info(f"💡 Model digunakan: {MODEL_NAME} | Device: {device_str.upper()}")

# Load model dan tokenizer
logging.info("📦 Memuat model dan tokenizer parafrasa...")
paraphrase_tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
paraphrase_model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME).to(device_str)
logging.info("✅ Model parafrasa berhasil dimuat.")

# Load model terjemahan
logging.info("🌐 Memuat model penerjemah ID↔EN...")
translator_indo_en = pipeline("translation", model="Helsinki-NLP/opus-mt-id-en", device=DEVICE)
translator_en_indo = pipeline("translation", model="Helsinki-NLP/opus-mt-en-id", device=DEVICE)
logging.info("✅ Model penerjemah berhasil dimuat.")

# ===============================
# 🧩 Fungsi bantu
# ===============================
def split_text_by_tokens(text, tokenizer, max_tokens=450):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, current = [], ""
    for sentence in sentences:
        if not sentence.strip():
            continue
        token_len = len(tokenizer.encode(current + " " + sentence))
        if token_len <= max_tokens:
            current += " " + sentence
        else:
            chunks.append(current.strip())
            current = sentence
    if current.strip():
        chunks.append(current.strip())
    return chunks

def translate_large_text_indo_en(text):
    try:
        start_time = time.time()
        logging.info("🌍 Memulai translasi ID→EN...")
        chunks = split_text_by_tokens(text, paraphrase_tokenizer)
        results = []
        for i, chunk in enumerate(chunks, start=1):
            logging.info(f"🔹 ID→EN chunk {i}/{len(chunks)} ({len(chunk.split())} kata)")
            translated = translator_indo_en(chunk)[0]["translation_text"]
            results.append(translated.strip())
        dur = time.time() - start_time
        logging.info(f"✅ Translasi ID→EN selesai dalam {dur:.2f} detik.")
        return " ".join(results)
    except Exception as e:
        logging.exception(f"⚠️ Gagal translate ID→EN: {e}")
        return text

def translate_large_text_en_indo(text):
    try:
        start_time = time.time()
        logging.info("🌍 Memulai translasi EN→ID...")
        chunks = split_text_by_tokens(text, paraphrase_tokenizer)
        results = []
        for i, chunk in enumerate(chunks, start=1):
            logging.info(f"🔹 EN→ID chunk {i}/{len(chunks)} ({len(chunk.split())} kata)")
            translated = translator_en_indo(chunk)[0]["translation_text"]
            results.append(translated.strip())
        dur = time.time() - start_time
        logging.info(f"✅ Translasi EN→ID selesai dalam {dur:.2f} detik.")
        return " ".join(results)
    except Exception as e:
        logging.exception(f"⚠️ Gagal translate EN→ID: {e}")
        return text

def paraphrase_indonesian_text(text: str) -> str:
    try:
        total_start = time.time()
        logging.info("🟢 Memulai proses parafrasa teks.")
        paragraphs = [p for p in text.split("\n") if p.strip()]
        all_results = []

        for p_idx, paragraph in enumerate(paragraphs, start=1):
            logging.info(f"📄 Paragraf {p_idx}/{len(paragraphs)} mulai diproses.")
            para_start = time.time()

            chunks = split_text_by_tokens(paragraph, paraphrase_tokenizer)
            para_results = []

            for c_idx, chunk in enumerate(chunks, start=1):
                chunk_start = time.time()
                logging.info(f"🧩 Chunk {c_idx}/{len(chunks)} | {len(chunk.split())} kata.")
                
                # 1️⃣ Translate ID → EN
                translated_en = translate_large_text_indo_en(chunk)

                # 2️⃣ Parafrase bahasa Inggris
                logging.info("✏️ Memulai proses parafrasa bahasa Inggris...")
                input_text = ( f"paraphrase this text in a formal and academic tone: {translated_en} </s>")

                with torch.no_grad():
                    inputs = paraphrase_tokenizer(
                        input_text,
                        return_tensors="pt",
                        truncation=True,
                        max_length=512
                    ).to(paraphrase_model.device)

                    outputs = paraphrase_model.generate(
                        **inputs,
                        max_length=2048,
                        min_length=100,
                        num_beams=8,
                        temperature=0.9,
                        top_p=0.92,
                        repetition_penalty=2.5,
                        early_stopping=False,
                        do_sample=True,
                        num_return_sequences=1
                    )
                    
                    # Log panjang token input/output
                    logging.info(f"✍️ Panjang token input: {len(inputs['input_ids'][0])}")
                    logging.info(f"📝 Panjang token output: {len(outputs[0])}")

                paraphrased_en = paraphrase_tokenizer.decode(outputs[0], skip_special_tokens=True)
                # 🔧 Bersihkan awalan yang tidak diinginkan
                paraphrased_en = re.sub(r"(?i)^paraphrasedoutput\s*[:\-]*\s*", "", paraphrased_en).strip()

                logging.info(f"✅ Parafrase EN selesai ({len(paraphrased_en.split())} kata).")

                # 3️⃣ Translate EN → ID
                translated_id = translate_large_text_en_indo(paraphrased_en)
                logging.info(f"✅ Chunk {c_idx} selesai total dalam {time.time() - chunk_start:.2f} detik.\n")
                para_results.append(translated_id.strip())

            all_results.append(" ".join(para_results))
            logging.info(f"✅ Paragraf {p_idx} selesai dalam {time.time() - para_start:.2f} detik.\n")

        logging.info(f"🎉 Semua paragraf selesai dalam {time.time() - total_start:.2f} detik total.")
        return "\n".join(all_results).strip()

    except Exception as e:
        logging.exception(f"❌ Gagal memproses teks panjang: {e}")
        return text

# ===============================
# ⚙️ Worker Asynchronous (Thread)
# ===============================
tasks = {}

def process_docx_in_background(task_id, file_path):
    try:
        document = Document(file_path)
        new_doc = Document()

        total_paras = len(document.paragraphs)
        logging.info(f"📘 Memproses dokumen DOCX ({total_paras} paragraf).")
        tasks[task_id] = {"status": "processing", "progress": 0, "total": total_paras}

        for i, para in enumerate(document.paragraphs, start=1):
            text = para.text.strip()
            logging.info(f"📄 Paragraf {i}/{total_paras} sedang diproses...")
            if not text:
                new_doc.add_paragraph("")
            else:
                rewritten = paraphrase_indonesian_text(text)
                new_doc.add_paragraph(rewritten)
            tasks[task_id]["progress"] = i

        output_path = f"../tmp_paraphrase/{task_id}.docx"
        new_doc.save(output_path)
        tasks[task_id]["status"] = "done"
        tasks[task_id]["output"] = output_path
        logging.info(f"✅ Dokumen selesai: {output_path}")
    except Exception as e:
        logging.exception(f"❌ Error pada task {task_id}: {e}")
        tasks[task_id] = {"status": "error", "error": str(e)}

# ===============================
# 📄 Endpoint - Parafrase TEXT
# ===============================
@app.route("/parafrase_text", methods=["POST"])
def paraphrase_text_api():
    text = request.form.get("text", "").strip()
    if not text:
        return jsonify({"error": "Teks tidak boleh kosong."}), 400

    logging.info("📨 Permintaan baru /parafrase_text diterima.")
    logging.info(f"🧾 Panjang input: {len(text.split())} kata.")
    result = paraphrase_indonesian_text(text)
    logging.info("🏁 Parafrase teks selesai dan hasil dikirim ke klien.")
    return jsonify({"result": result}), 200

# ===============================
# 📄 Endpoint - Parafrase DOCX Async
# ===============================
@app.route("/parafrase_docx", methods=["POST"])
def paraphrase_docx():
    file = request.files.get("file")
    if not file or not file.filename.endswith(".docx"):
        return jsonify({"error": "File DOCX tidak valid."}), 400

    tmp_path = f"../tmp_paraphrase/{uuid.uuid4()}.docx"
    file.save(tmp_path)
    logging.info(f"📁 File DOCX diterima dan disimpan sementara di {tmp_path}")

    task_id = str(uuid.uuid4())
    thread = threading.Thread(target=process_docx_in_background, args=(task_id, tmp_path))
    thread.start()

    logging.info(f"🚀 Task {task_id} dimulai untuk file DOCX.")
    return jsonify({"task_id": task_id, "message": "Dokumen sedang diproses di latar belakang."}), 202

# ===============================
# 📊 Endpoint - Cek Progress
# ===============================
@app.route("/parafrase_status/<task_id>", methods=["GET"])
def check_status(task_id):
    task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "Task tidak ditemukan."}), 404
    return jsonify(task), 200

# ===============================
# 📄 Endpoint - Unduh Hasil
# ===============================
@app.route("/parafrase_download/<task_id>", methods=["GET"])
def download_paraphrased_docx(task_id):
    task = tasks.get(task_id)
    if not task or task.get("status") != "done":
        return jsonify({"error": "File belum tersedia atau gagal diproses."}), 400
    return send_file(task["output"], as_attachment=True, download_name="Parafrase_Selesai.docx")

# ===============================
# 🚀 Jalankan Server
# ===============================
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True, threaded=True)